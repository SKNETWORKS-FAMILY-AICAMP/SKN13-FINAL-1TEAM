from fastapi import FastAPI, Depends, APIRouter, HTTPException, File, UploadFile, Request
from fastapi.responses import Response, StreamingResponse, JSONResponse, FileResponse
from starlette.background import BackgroundTask
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Generator, Optional
import json, uuid, shutil, tempfile, os
from pathlib import Path
from sqlalchemy.orm import Session
from datetime import datetime
import fitz # PyMuPDF
from docx import Document as DocxDocument # python-docx
from langchain_core.messages.tool import ToolMessage
import time
from .ChatBot.agents.RoutingAgent import RoutingAgent, generate_config
from .ChatBot.core.AgentState import AgentState
from .database import create_db_and_tables, SessionLocal, ChatSession, ChatMessage, ToolMessageRecord, User, Calendar, Event, Document
from .ChatBot.tools.html_to_docx import convert_html_to_docx

# FastAPI 인스턴스 생성
app = FastAPI()

# API 라우터 생성
api_router = APIRouter()

# CORS 미들웨어 설정
origins = [
    "http://localhost",
    "http://localhost:5173",  # 프론트엔드 개발 서버
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Document Storage Paths ---
UPLOAD_DIR = Path("uploaded_files")
EDITABLE_MD_DIR = Path("editable_markdown")

# 데이터베이스 테이블 생성 및 디렉토리 생성
@app.on_event("startup")
def on_startup():
    create_db_and_tables()
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    EDITABLE_MD_DIR.mkdir(parents=True, exist_ok=True)

# DB 세션 의존성
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# --- Chat Message Handling ---
def _create_chat_message(db: Session, session_id: str, role: str, content: str, message_id: str = None):
    message = ChatMessage(
        session_id=session_id,
        role=role,
        content=content,
        message_id=message_id,
        timestamp=datetime.now()
    )
    db.add(message)
    db.commit()
    db.refresh(message)
    return message

def _create_tool_message(
    db: Session,
    chat_msg_id: int,
    tool_call_id: str | None = None,
    status: str | None = None,
    artifact: dict | None = None,
    raw_content: dict | str | None = None,
) -> ToolMessageRecord:
    """ToolMessageRecord를 DB에 저장하고 반환"""
    tool_record = ToolMessageRecord(
        chat_message_id=chat_msg_id,
        tool_call_id=tool_call_id,
        tool_status=status,
        tool_artifact=artifact,
        tool_raw_content=raw_content,
    )
    db.add(tool_record)
    db.commit()
    db.refresh(tool_record)
    return tool_record


async def _handle_tool_start(event: dict, session_id: str, db: Session):
    tool_name = event.get("name", "Unknown Tool")
    tool_input = event.get("data", {}).get("input", {})
    tool_call_id = event.get("tool_call_id")
    tool_artifact = event.get("artifact")

    # 사용자 화면용 메시지
    thinking_message = (
        f"[AI Thinking]: Using tool '{tool_name}' with input:\n"
        f"```json\n{json.dumps(tool_input, indent=2, ensure_ascii=False)}\n```"
    )

    # 1. ChatMessage 저장
    chat_msg = _create_chat_message(db, session_id, "assistant", thinking_message)

    # 2. ToolMessageRecord 저장
    tool_raw_content = {"tool_name": tool_name, "input": tool_input}
    _create_tool_message(
        db,
        chat_msg.id,
        tool_call_id=tool_call_id,
        status="started",
        artifact=tool_artifact,
        raw_content=tool_raw_content,
    )

    # 3. SSE 전송
    yield f"data: {json.dumps({'thinking_message': thinking_message}, ensure_ascii=False)}\n\n"


async def _handle_tool_end(event: dict, session_id: str, db: Session):
    tool_name = event.get("name")
    raw_output = event.get("data", {}).get("output", "")
    print(f"--- _handle_tool_end called. tool_name: {tool_name}, raw_output type: {type(raw_output)} ---")

    # If the editor tool finishes, its output is the new document.
    # Yield a specific event for the frontend to catch and update the editor.
    DOCUMENT_UPDATE_TOOLS = {
    "run_document_edit",  # 메인 편집 도구 (기존)
    "replace_text_in_document",  # 기존 도구
    "create_document_structure",  # 문서 구조 생성
    "create_business_report_template",  # 보고서 템플릿
    "create_meeting_minutes_template",  # 회의록 템플릿
    "add_formatted_text",  # 텍스트 스타일링
    "create_list",  # 리스트 생성
    "create_table",  # 테이블 생성
    "add_blockquote",  # 인용문/들여쓰기
    "format_text_block",  # 텍스트 블록 포맷팅
    "apply_document_styling",  # 스타일 적용
    "enhance_document_readability",  # 가독성 향상
    }
    if tool_name in DOCUMENT_UPDATE_TOOLS:
        content_to_send = raw_output.content if isinstance(raw_output, ToolMessage) else raw_output
        print(f"--- Sending document_update for replace_text_in_document. Content length: {len(content_to_send) if isinstance(content_to_send, str) else 'N/A'} ---")
        yield f"data: {json.dumps({'document_update': content_to_send}, ensure_ascii=False)}\n\n"

    formatted_output = "[Tool Output]: "
    tool_raw_json = None

    try:
        if isinstance(raw_output, ToolMessage):
            tool_raw_json = {
                "content": raw_output.content,
                "type": getattr(raw_output, "type", None),
                "tool_call_id": getattr(raw_output, "tool_call_id", None),
                "artifact": getattr(raw_output, "artifact", None),
                "status": getattr(raw_output, "status", None),
            }
            parsed_output = raw_output.content
        else:
            parsed_output = str(raw_output)
            tool_raw_json = {"raw": parsed_output}

        # 사용자용 포맷팅
        if isinstance(parsed_output, str):
            try:
                parsed_json = json.loads(parsed_output)
                formatted_output += (
                    f"\n```json\n{json.dumps(parsed_json, indent=2, ensure_ascii=False)}\n```"
                )
            except json.JSONDecodeError:
                formatted_output += f"\n```\n{parsed_output}\n```"
        elif isinstance(parsed_output, dict):
            formatted_output += (
                f"\n```json\n{json.dumps(parsed_output, indent=2, ensure_ascii=False)}\n```"
            )
        else:
            formatted_output += f"`{str(parsed_output)}`"

    except Exception as e:
        formatted_output += f"`Error processing output: {e}`"
        print(f"[Error] raw_output={raw_output} -> {e}")

    # 1. SSE 전송
    yield f"data: {json.dumps({'tool_message': formatted_output}, ensure_ascii=False)}\n\n"

    # 2. ChatMessage 저장
    chat_msg = _create_chat_message(
        db,
        session_id,
        "tool",
        content=json.dumps(tool_raw_json, ensure_ascii=False),
    )

    # 3. ToolMessageRecord 저장
    _create_tool_message(
        db,
        chat_msg.id,
        tool_call_id=tool_raw_json.get("tool_call_id"),
        status=tool_raw_json.get("status", "finished"),
        artifact=tool_raw_json.get("artifact"),
        raw_content=tool_raw_json,
    )


# --- Pydantic Models (Existing) ---
class MessageSaveRequest(BaseModel):
    session_id: str
    role: str
    content: str

class EventBase(BaseModel):
    title: str
    description: Optional[str] = None
    start: datetime
    end: Optional[datetime] = None
    all_day: bool = False
    color: Optional[str] = None

class EventCreate(EventBase):
    pass

class EventUpdate(EventBase):
    pass

class EventInDB(EventBase):
    id: int
    calendar_id: int

    class Config:
        orm_mode = True

class EventOut(BaseModel):
    id: int
    title: str
    description: Optional[str] = None
    start: str
    end: Optional[str] = None
    allDay: bool
    color: Optional[str] = None

class ExportDocxRequest(BaseModel):
    html: str
    filename: str = "exported_document.docx"

# --- Helper for Document Conversion ---
def _convert_to_markdown(file_path: Path, file_type: str) -> str:
    content = ""
    try:
        if file_type == "pdf":
            doc = fitz.open(file_path)
            for page in doc:
                content += page.get_text()
            doc.close()
        elif file_type == "docx":
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif file_type in ["hwp", "hwpx"]:
            # Assuming read_hwpx_file_content returns text content
            content = None#read_hwpx_file_content(str(file_path))
        elif file_type in ["md", "txt"]:
            content = file_path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type for conversion: {file_type}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File conversion failed: {e}")
    return content

# --- Document Management API Endpoints ---

@api_router.post("/documents/upload")
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    file_extension = Path(file.filename).suffix.lower().lstrip('.')
    if file_extension not in ["pdf", "docx", "hwp", "hwpx", "md", "txt"]:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")

    doc_id = str(uuid.uuid4())
    original_file_path = UPLOAD_DIR / f"{doc_id}_{file.filename}"
    markdown_file_path = EDITABLE_MD_DIR / f"{doc_id}.md"

    # Save original file
    try:
        with open(original_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not save original file: {e}")

    # Convert to Markdown
    try:
        markdown_content = _convert_to_markdown(original_file_path, file_extension)
        with open(markdown_file_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
    except Exception as e:
        # Clean up if conversion fails
        original_file_path.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail=f"File conversion to Markdown failed: {e}")

    # Save metadata to DB
    db_document = Document(
        id=doc_id,
        original_filename=file.filename,
        file_type=file_extension,
        original_file_path=str(original_file_path),
        markdown_file_path=str(markdown_file_path)
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)

    return JSONResponse(content={"doc_id": doc_id, "markdown_content": markdown_content, "message": "Document uploaded and converted successfully"})

@api_router.post("/documents/export/docx")
async def export_document_as_docx(req: ExportDocxRequest, request: Request):
    print("--- Raw body ---")
    raw_body = await request.body()
    print(raw_body.decode())

    print(f"--- Parsed html_content length: {len(req.html)} ---")
    print(f"--- Parsed filename: {req.filename} ---")

    safe_filename = req.filename.strip()
    if not safe_filename.endswith(".docx"):
        safe_filename += ".docx"
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_filepath = temp_file.name

    doc_title = os.path.splitext(safe_filename)[0]
    success = convert_html_to_docx(req.html, temp_filepath, title=doc_title)

    if not success:
        raise HTTPException(status_code=500, detail="Failed to convert HTML to DOCX.")
    
    return FileResponse(
        path=temp_filepath,
        filename=safe_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(os.unlink, temp_filepath)
    )

@api_router.get("/documents/{doc_id}/content")
async def get_document_content(doc_id: str, db: Session = Depends(get_db)):
    db_document = db.query(Document).filter(Document.id == doc_id).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    markdown_file_path = Path(db_document.markdown_file_path)
    if not markdown_file_path.exists():
        raise HTTPException(status_code=404, detail="Markdown content not found for this document")
    
    content = markdown_file_path.read_text(encoding="utf-8")
    return JSONResponse(content={"doc_id": doc_id, "markdown_content": content})

@api_router.put("/documents/{doc_id}/save_content")
async def save_document_content(doc_id: str, content: str, db: Session = Depends(get_db)):
    db_document = db.query(Document).filter(Document.id == doc_id).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    markdown_file_path = Path(db_document.markdown_file_path)
    
    try:
        markdown_file_path.write_text(content, encoding="utf-8")
        db_document.updated_at = datetime.now()
        db.commit()
        db.refresh(db_document)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save markdown content: {e}")
    
    return JSONResponse(content={"doc_id": doc_id, "message": "Markdown content saved successfully"})

@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, db: Session = Depends(get_db)):
    db_document = db.query(Document).filter(Document.id == doc_id).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    original_file_path = Path(db_document.original_file_path)
    markdown_file_path = Path(db_document.markdown_file_path)

    try:
        if original_file_path.exists():
            original_file_path.unlink()
        if markdown_file_path.exists():
            markdown_file_path.unlink()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete files: {e}")
    
    db.delete(db_document)
    db.commit()
    
    return JSONResponse(content={"message": "Document and associated files deleted successfully"})

# --- Existing API Endpoints (Calendar, Chat) ---

# Helper function to get or create a default calendar for a user
def get_or_create_default_calendar(db: Session, user_id: int) -> Calendar:
    calendar = db.query(Calendar).filter(Calendar.user_id == user_id).first()
    if not calendar:
        calendar = Calendar(user_id=user_id, name="My Calendar")
        db.add(calendar)
        db.commit()
        db.refresh(calendar)
    return calendar

@api_router.post("/events", response_model=EventOut)
def create_event(event: EventCreate, db: Session = Depends(get_db)):
    default_user_id = 1  # 임시 사용자
    calendar = get_or_create_default_calendar(db, default_user_id)
    
    db_event = Event(
        **event.dict(),
        calendar_id=calendar.id
    )
    db.add(db_event)
    db.commit()
    db.refresh(db_event)
    
    return EventOut(
        id=db_event.id,
        title=db_event.title,
        description=db_event.description,
        start=db_event.start.isoformat(),
        end=db_event.end.isoformat() if db_event.end else None,
        allDay=db_event.all_day,
        color=db_event.color
    )

@api_router.get("/events", response_model=List[EventOut])
def get_events(start: datetime, end: datetime, db: Session = Depends(get_db)):
    default_user_id = 1 # 임시 사용자
    
    calendar = get_or_create_default_calendar(db, default_user_id)
    
    events = db.query(Event).filter(
        Event.calendar_id == calendar.id,
        Event.start < end,
        Event.end > start
    ).all()
    
    return [
        EventOut(
            id=e.id,
            title=e.title,
            description=e.description,
            start=e.start.isoformat(),
            end=e.end.isoformat() if e.end else None,
            allDay=e.all_day,
            color=e.color
        ) for e in events
    ]

@api_router.put("/events/{event_id}", response_model=EventOut)
def update_event(event_id: int, event: EventUpdate, db: Session = Depends(get_db)):
    db_event = db.query(Event).filter(Event.id == event_id).first()
    if not db_event:
        raise HTTPException(status_code=404, detail="Event not found")

    default_user_id = 1
    if db_event.calendar.user_id != default_user_id:
        raise HTTPException(status_code=403, detail="Not authorized to update this event")

    for key, value in event.dict().items():
        setattr(db_event, key, value)
    
    db.commit()
    db.refresh(db_event)

    return EventOut(
        id=db_event.id,
        title=db_event.title,
        description=db_event.description,
        start=db_event.start.isoformat(),
        end=db_event.end.isoformat() if db_event.end else None,
        allDay=db_event.all_day,
        color=db_event.color
    )

@api_router.delete("/events/{event_id}", status_code=204)
def delete_event(event_id: int, db: Session = Depends(get_db)):
    db_event = db.query(Event).filter(Event.id == event_id).first()
    if not db_event:
        raise HTTPException(status_code=404, detail="Event not found")

    default_user_id = 1
    if db_event.calendar.user_id != default_user_id:
        raise HTTPException(status_code=403, detail="Not authorized to delete this event")

    db.delete(db_event)
    db.commit()
    return

@api_router.post("/chat/save")
async def save_message(request: MessageSaveRequest, db: Session = Depends(get_db)):
    default_user_id = 1
    user = db.query(User).filter(User.id == default_user_id).first()
    if not user:
        user = User(id=default_user_id, unique_auth_number="default_auth", username="default_user", hashed_password="", email="default@example.com")
        db.add(user)
        db.commit()
        db.refresh(user)

    session = db.query(ChatSession).filter(ChatSession.id == request.session_id).first()
    if not session:
        session = ChatSession(id=request.session_id, user_id=user.id, title="새로운 대화")
        db.add(session)
        db.commit()
        db.refresh(session)

    new_message = ChatMessage(session_id=request.session_id, role=request.role, content=request.content)
    db.add(new_message)
    db.commit()
    db.refresh(new_message)

    if request.role == "user" and session.title == "새로운 대화":
        session.title = request.content[:50]
        db.commit()

    return {"status": "success"}

@api_router.get("/chat/sessions")
async def get_chat_sessions(db: Session = Depends(get_db)):
    default_user_id = 1
    sessions = db.query(ChatSession).filter(ChatSession.user_id == default_user_id).order_by(ChatSession.created_at.desc()).all()
    sessions_list = []
    for session in sessions:
        sessions_list.append({"id": session.id, "title": session.title})
    return {"sessions": sessions_list}

@api_router.get("/chat/messages/{session_id}")
async def get_messages(session_id: str, db: Session = Depends(get_db)):
    messages = db.query(ChatMessage).filter(ChatMessage.session_id == session_id).order_by(ChatMessage.timestamp).all()
    return {"messages": [{"role": msg.role, "content": msg.content} for msg in messages]}

@api_router.get("/llm/stream")
async def llm_stream(session_id: str, prompt: str, document_content: Optional[str] = None, db: Session = Depends(get_db)):
    config = generate_config(session_id)
    chat_agent = RoutingAgent()

    return StreamingResponse(_stream_llm_response(session_id, prompt, document_content, chat_agent, config, db), media_type="text/event-stream")

async def _stream_llm_response(session_id: str, prompt: str, document_content: Optional[str], chat_agent, config, db: Session) -> Generator:
    # Get or create the chat session to prevent foreign key violations
    session = db.query(ChatSession).filter(ChatSession.id == session_id).first()
    if not session:
        default_user_id = 1
        user = db.query(User).filter(User.id == default_user_id).first()
        if not user:
            user = User(id=default_user_id, unique_auth_number="default_auth", username="default_user", hashed_password="", email="default@example.com")
            db.add(user)
        session = ChatSession(id=session_id, user_id=default_user_id, title="새로운 대화")
        db.add(session)
        db.commit()

    full_response_content = ""
    llm_output_buffer = "" # 라우팅 LLM 출력을 위한 버퍼
    
    history_messages = db.query(ChatMessage).filter(ChatMessage.session_id == session_id).order_by(ChatMessage.timestamp).all()
    
    messages = []
    for msg in history_messages:
        if msg.role == "tool":
            continue
        role = "ai" if msg.role == "assistant" else msg.role
        messages.append((role, msg.content))
    
    if not messages or messages[-1] != ("user", prompt):
        messages.append(("user", prompt))

    # ✅ Create the initial state object
    initial_state: AgentState = {
        "prompt": prompt,
        "document_content": document_content,
        "messages": messages,
        # The following fields will be populated by the agents
        "intent": None,
        "needs_document_content": False,
        "intermediate_steps": [],
        "generation": None,
    }

    # Pass the state object to the agent stream
    async for event in chat_agent.astream_events(initial_state, config=config):
        
        kind = event["event"]
        name = event.get("name")

        # 'request_document' 노드가 종료되었는지 확인
        if kind == "on_chain_end" and name == "request_document":
            node_output = event.get("data", {}).get("output", {})
            if node_output and node_output.get("needs_document_content"):
                print("--- 프론트엔드에 문서 요청 신호 전송 ---")
                tool_call_id = f"req_doc_{uuid.uuid4()}"
                yield f"data: {json.dumps({'needs_document_content': True, 'agent_context': {'tool_call_id': tool_call_id}}, ensure_ascii=False)}\n\n"
                

        # 'route_question' 체인이 종료되었고, 'request_document'가 트리거되지 않았다면 버퍼된 LLM 출력을 yield
        if kind == "on_chain_end" and name == "route_question":
            if llm_output_buffer:
                yield f"data: {json.dumps({'content': llm_output_buffer}, ensure_ascii=False)}\n\n"
                llm_output_buffer = "" # 버퍼 비우기

        # 일반적인 LLM 스트림 처리 (라우팅 LLM 제외)
        if kind == "on_chat_model_stream":
            content = event["data"]["chunk"].content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
                full_response_content += content
                
        elif kind == "on_tool_start":
            async for chunk in _handle_tool_start(event, session_id, db):
                yield chunk
        
        elif kind == "on_tool_end":
            async for chunk in _handle_tool_end(event, session_id, db):
                yield chunk
        
        elif kind == "on_end":
            final_state = event.get("data", {}).get("output", {})
            # ... (rest of the logic)
            yield "data: [DONE]\n\n"

    if full_response_content:
        _create_chat_message(db, session_id, "assistant", full_response_content)

@api_router.get("/open")
async def open_document(url):
    pass


@api_router.post("/logout")
async def logout(response: Response):
    response.delete_cookie(key="refresh_token")
    return {"message": "Logout successful"}


# API 라우터를 앱에 포함
app.include_router(api_router, prefix="/api/v1")
