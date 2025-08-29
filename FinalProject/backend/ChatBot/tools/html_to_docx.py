import os
import re
from typing import Dict, Union, Optional
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


class HTMLToDocxConverter:
    """HTML을 DOCX 형식으로 변환하는 클래스"""
    
    def __init__(self):
        self.alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
    
    def parse_css_style(self, style_str: str) -> Dict[str, Union[str, Dict[str, int]]]:
        """CSS 스타일 문자열을 파싱하여 딕셔너리로 반환"""
        style_dict = {}
        if not style_str:
            return style_dict

        for item in style_str.split(';'):
            if ':' not in item:
                continue
            
            key, value = item.split(':', 1)
            key = key.strip().lower()
            value = value.strip()
            
            if key == 'font-family':
                style_dict['font_family'] = self._clean_font_family(value)
            elif key == 'font-size':
                style_dict['font_size'] = value
            elif key == 'color':
                style_dict['color'] = self._parse_color(value)
            elif key == 'background-color':
                style_dict['highlight'] = value
            elif key == 'text-align':
                style_dict['text_align'] = value
                
        return style_dict
    
    def _clean_font_family(self, value: str) -> str:
        """폰트 패밀리 값에서 따옴표 제거"""
        return value.replace('"', '').replace("'", "")
    
    def _parse_color(self, value: str) -> Union[str, Dict[str, int]]:
        """색상 값을 파싱 (RGB 또는 색상명)"""
        rgb_match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', value)
        if rgb_match:
            return {
                'r': int(rgb_match.group(1)),
                'g': int(rgb_match.group(2)),
                'b': int(rgb_match.group(3))
            }
        return value
    
    def apply_text_formatting(self, run, style: Dict):
        """텍스트 런에 포맷팅 적용"""
        font = run.font
        
        # 기본 텍스트 스타일 적용
        if style.get('bold'):
            font.bold = True
        if style.get('italic'):
            font.italic = True
        if style.get('underline'):
            font.underline = True
        if style.get('strike'):
            font.strike = True
            
        # 폰트 패밀리 설정
        if style.get('font_family'):
            font.name = style['font_family']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_family'])
            
        # 폰트 크기 설정
        if style.get('font_size'):
            try:
                size_pt = float(re.sub(r'[^\d.]', '', style['font_size']))
                font.size = Pt(size_pt)
            except ValueError:
                pass
                
        # 글자 색상 설정
        if style.get('color') and isinstance(style['color'], dict):
            color = style['color']
            font.color.rgb = RGBColor(color['r'], color['g'], color['b'])
            
        # 하이라이트 설정
        if style.get('highlight'):
            font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    def create_paragraph_container(self, container) -> Paragraph:
        """컨테이너에 새로운 문단 생성"""
        if isinstance(container, (Document, _Cell)):
            return container.add_paragraph()
        elif isinstance(container, Paragraph):
            return container
        else:
            raise TypeError(f"지원하지 않는 컨테이너 타입: {type(container)}")
    
    def process_text_node(self, node: NavigableString, container, style: Dict):
        """텍스트 노드 처리"""
        text = str(node).strip()
        if not text:
            return
            
        if isinstance(container, Paragraph):
            run = container.add_run(text)
            self.apply_text_formatting(run, style)
        elif isinstance(container, (_Cell, Document)):
            paragraph = container.add_paragraph()
            run = paragraph.add_run(text)
            self.apply_text_formatting(run, style)
    
    def process_block_element(self, node: Tag, container, style: Dict):
        """블록 레벨 요소 처리 (p, h1-h6)"""
        tag = node.name
        
        if tag.startswith('h') and len(tag) == 2 and tag[1].isdigit():
            level = int(tag[1])
            paragraph = self.create_paragraph_container(container)
            paragraph.style = f'Heading {level}'
        else:
            paragraph = self.create_paragraph_container(container)
        
        # 텍스트 정렬 설정
        if style.get('text_align'):
            paragraph.alignment = self.alignment_map.get(
                style['text_align'], 
                WD_ALIGN_PARAGRAPH.LEFT
            )
        
        # 자식 노드들 처리
        for child in node.children:
            self.process_node(child, paragraph, style)
    
    def process_list_element(self, node: Tag, container, style: Dict):
        """리스트 요소 처리 (ul, ol)"""
        list_items = node.find_all('li', recursive=False)
        style_name = 'List Bullet' if node.name == 'ul' else 'List Number'
        
        for li in list_items:
            paragraph = self.create_paragraph_container(container)
            paragraph.style = style_name
            
            for child in li.children:
                self.process_node(child, paragraph, style)
    
    def process_table_element(self, node: Tag, container, style: Dict):
        """테이블 요소 처리"""
        rows = node.find_all('tr')
        if not rows:
            return
            
        # 테이블 크기 계산
        col_count = max(len(tr.find_all(['td', 'th'])) for tr in rows)
        
        # 테이블 생성
        if isinstance(container, Paragraph):
            table = container._parent.add_table(rows=len(rows), cols=col_count, style='Table Grid')
        else:
            table = container.add_table(rows=len(rows), cols=col_count, style='Table Grid')
        
        # 테이블 내용 채우기
        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(['td', 'th'])
            for col_idx, cell_node in enumerate(cells):
                if col_idx < col_count:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''  # 기본 텍스트 제거
                    
                    for child in cell_node.children:
                        self.process_node(child, cell, style)
    
    def process_inline_element(self, node: Tag, container, style: Dict):
        """인라인 요소 처리"""
        for child in node.children:
            self.process_node(child, container, style)
    
    def update_style_from_node(self, node: Tag, base_style: Dict) -> Dict:
        """HTML 노드에서 스타일 정보 추출하여 기존 스타일에 병합"""
        new_style = base_style.copy()
        tag = node.name
        
        # 태그별 기본 스타일 설정
        tag_styles = {
            'b': {'bold': True}, 'strong': {'bold': True},
            'i': {'italic': True}, 'em': {'italic': True},
            'u': {'underline': True},
            's': {'strike': True}, 'del': {'strike': True},
            'mark': {'highlight': 'yellow'}
        }
        
        if tag in tag_styles:
            new_style.update(tag_styles[tag])
        
        # CSS 스타일 속성 적용
        if node.has_attr('style'):
            css_style = self.parse_css_style(node['style'])
            new_style.update(css_style)
            
        return new_style
    
    def process_node(self, node: Union[Tag, NavigableString], container, style: Dict):
        """노드를 재귀적으로 처리하는 메인 함수"""
        # 텍스트 노드 처리
        if isinstance(node, NavigableString):
            self.process_text_node(node, container, style)
            return
        
        # 태그 노드가 아닌 경우 무시
        if not isinstance(node, Tag):
            return
            
        # 스타일 업데이트
        updated_style = self.update_style_from_node(node, style)
        tag = node.name
        
        # 태그별 처리
        if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.process_block_element(node, container, updated_style)
        elif tag in ['ul', 'ol']:
            self.process_list_element(node, container, updated_style)
        elif tag == 'table':
            self.process_table_element(node, container, updated_style)
        else:
            # 기타 인라인 요소들
            self.process_inline_element(node, container, updated_style)
    
    def convert(self, html_content: str, output_path: str, title: str = "") -> bool:
        """HTML을 DOCX로 변환하는 메인 메서드"""
        if not output_path.endswith('.docx'):
            print(f"오류: 출력 파일은 .docx 확장자여야 합니다.")
            return False
            
        try:
            # 새 문서 생성
            doc = Document()
            
            # 제목 추가 (옵션)
            if title:
                doc.add_heading(title, level=0)
            
            # HTML 파싱
            soup = BeautifulSoup(html_content, 'html.parser')
            body = soup.find('body') or soup
            
            # 각 최상위 요소 처리
            for element in body.children:
                self.process_node(element, doc, {})
            
            # 문서 저장
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"변환 중 오류 발생: {e}")
            import traceback
            traceback.print_exc()
            return False


def convert_html_to_docx(html_content: str, output_path: str, title: str = "") -> bool:
    """편의를 위한 함수형 인터페이스"""
    converter = HTMLToDocxConverter()
    return converter.convert(html_content, output_path, title)


# ------------------------ 사용 예시 ------------------------
def main():
    """메인 실행 함수"""
    sample_html = '''
    <h1>H1 제목</h1>
    <h2>H2 제목</h2>
    <p><span style="font-family: 돋움;">돋움 글씨</span></p>
    <p><span style="font-size: 48px;">48px 글씨</span></p>
    <p><strong>볼드체</strong></p>
    <p><em>이탤릭</em></p>
    <p><u>밑줄</u></p>
    <p><s>취소선</s></p>
    <p><span style="color: rgb(255,0,0);">빨간 글씨</span></p>
    <p><mark>하이라이트</mark></p>
    <table>
        <tr><td>1</td><td>2</td></tr>
        <tr><td>3</td><td>4</td></tr>
    </table>
    <ol>
        <li>번호 1</li>
        <li>번호 2</li>
    </ol>
    <ul>
        <li>글머리 1</li>
        <li>글머리 2</li>
    </ul>
    '''
    
    # 출력 경로 설정
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    output_file = os.path.join(desktop_path, "refactored_doc.docx")
    
    # 변환 실행
    success = convert_html_to_docx(
        html_content=sample_html,
        output_path=output_file,
        title="리팩토링된 테스트 문서"
    )
    
    print("변환 완료!" if success else "변환 실패!")


if __name__ == '__main__':
    main()