from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Request, Body
from fastapi.responses import JSONResponse, FileResponse
from starlette.background import BackgroundTask
from pydantic import BaseModel
from typing import Optional, List
import json, uuid, shutil, tempfile, os
from pathlib import Path
from sqlalchemy.orm import Session
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document as DocxDocument  # python-docx

from ..database import get_db, Document, User
from ..ChatBot.tools.html_to_docx import convert_html_to_docx
from .auth_routes import get_current_user

# Presigned URL 도구 import
try:
    from ..ChatBot.tools.presigned import (
        get_upload_url, get_download_url, get_public_url,
        upload_file_directly, OneTimePresignedURLManager
    )
    PRESIGNED_AVAILABLE = True
except ImportError:
    PRESIGNED_AVAILABLE = False

# APIRouter 인스턴스 생성
router = APIRouter()

# --- 문서 저장 경로 ---
UPLOAD_DIR = Path("uploaded_files")
EDITABLE_MD_DIR = Path("editable_markdown")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
EDITABLE_MD_DIR.mkdir(parents=True, exist_ok=True)

# --- Pydantic 모델 ---
class ExportDocxRequest(BaseModel):
    html: str
    filename: str = "exported_document.docx"

class DocumentOut(BaseModel):
    id: str
    original_filename: str
    file_type: str
    uploaded_at: str
    updated_at: Optional[str]

    class Config:
        orm_mode = True

class PresignedURLRequest(BaseModel):
    filename: str
    contentType: Optional[str] = "application/octet-stream"

# --- 변환 헬퍼 ---
def _convert_to_markdown(file_path: Path, file_type: str) -> str:
    content = ""
    try:
        if file_type == "pdf":
            doc = fitz.open(file_path)
            for page in doc:
                content += page.get_text()
            doc.close()
        elif file_type == "docx":
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif file_type in ["hwp", "hwpx"]:
            # TODO: hwp/hwpx 변환기 연결
            content = ""
        elif file_type in ["md", "txt"]:
            content = file_path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type for conversion: {file_type}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File conversion failed: {e}")
    return content

# --- 문서 관리 API ---

@router.get("/", response_model=List[DocumentOut])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 보안: 자신의 문서만 조회 가능
    documents = db.query(Document).filter(Document.owner_id == current_user.id).all()
    return [
        DocumentOut(
            id=doc.id,
            original_filename=doc.original_filename,
            file_type=doc.file_type,
            uploaded_at=doc.uploaded_at.isoformat(),
            updated_at=doc.updated_at.isoformat() if doc.updated_at else None,
        )
        for doc in documents
    ]

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    file_extension = Path(file.filename).suffix.lower().lstrip('.')
    if file_extension not in ["pdf", "docx", "hwp", "hwpx", "md", "txt"]:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")

    safe_filename = Path(file.filename).name
    doc_id = str(uuid.uuid4())
    original_file_path = UPLOAD_DIR / f"{doc_id}_{safe_filename}"
    markdown_file_path = EDITABLE_MD_DIR / f"{doc_id}.md"

    try:
        with open(original_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not save original file: {e}")

    try:
        markdown_content = _convert_to_markdown(original_file_path, file_extension)
        with open(markdown_file_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
    except Exception as e:
        original_file_path.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail=f"File conversion to Markdown failed: {e}")

    db_document = Document(
        id=doc_id,
        owner_id=current_user.id,  # 소유자 설정
        original_filename=safe_filename,
        file_type=file_extension,
        original_file_path=str(original_file_path),
        markdown_file_path=str(markdown_file_path),
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)

    return JSONResponse(content={
        "doc_id": doc_id,
        "markdown_content": markdown_content,
        "message": "Document uploaded and converted successfully"
    })

@router.post("/{doc_id}/export")
async def export_document_as_docx(
    doc_id: str, 
    req: ExportDocxRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 보안: 소유자만 내보내기 가능
    db_document = db.query(Document).filter(
        Document.id == doc_id,
        Document.owner_id == current_user.id
    ).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    safe_filename = req.filename.strip()
    if not safe_filename.endswith(".docx"):
        safe_filename += ".docx"

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_filepath = temp_file.name

    doc_title = os.path.splitext(safe_filename)[0]
    success = convert_html_to_docx(req.html, temp_filepath, title=doc_title)

    if not success:
        raise HTTPException(status_code=500, detail="Failed to convert HTML to DOCX.")

    return FileResponse(
        path=temp_filepath,
        filename=safe_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(os.unlink, temp_filepath)
    )

@router.get("/{doc_id}/content")
async def get_document_content(
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 보안: 소유자만 접근 가능
    db_document = db.query(Document).filter(
        Document.id == doc_id,
        Document.owner_id == current_user.id
    ).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")

    markdown_file_path = Path(db_document.markdown_file_path)
    if not markdown_file_path.exists():
        raise HTTPException(status_code=404, detail="Markdown content not found for this document")

    content = markdown_file_path.read_text(encoding="utf-8")
    return JSONResponse(content={"doc_id": doc_id, "markdown_content": content})

@router.put("/{doc_id}/content")
async def save_document_content(
    doc_id: str, 
    content: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 보안: 소유자만 수정 가능
    db_document = db.query(Document).filter(
        Document.id == doc_id,
        Document.owner_id == current_user.id
    ).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")

    markdown_file_path = Path(db_document.markdown_file_path)

    try:
        markdown_file_path.write_text(content, encoding="utf-8")
        db_document.updated_at = datetime.now()
        db.commit()
        db.refresh(db_document)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save markdown content: {e}")

    return JSONResponse(content={"doc_id": doc_id, "message": "Markdown content saved successfully"})

@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 보안: 소유자만 삭제 가능
    db_document = db.query(Document).filter(
        Document.id == doc_id,
        Document.owner_id == current_user.id
    ).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")

    original_file_path = Path(db_document.original_file_path)
    markdown_file_path = Path(db_document.markdown_file_path)

    try:
        if original_file_path.exists():
            original_file_path.unlink()
        if markdown_file_path.exists():
            markdown_file_path.unlink()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete files: {e}")

    db.delete(db_document)
    db.commit()

    return JSONResponse(content={"message": "Document and associated files deleted successfully"})

# --- Presigned URL 엔드포인트 (main2.py에서 이식) ---

@router.post("/presigned")
async def create_presigned_url(
    request: PresignedURLRequest,
    current_user: User = Depends(get_current_user)
):
    """
    S3 업로드를 위한 presigned URL 발급
    프론트엔드가 직접 S3에 파일을 업로드할 수 있도록 함
    """
    if not PRESIGNED_AVAILABLE:
        raise HTTPException(
            status_code=503, 
            detail="Presigned URL service not available. Check S3 configuration."
        )
    
    try:
        # main2.py의 로직 복사
        result = get_upload_url(request.filename, request.contentType)
        return {
            "uploadUrl": result.get("uploadUrl"),
            "fileKey": result.get("fileKey"), 
            "displayName": request.filename
        }
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"Failed to generate presigned URL: {str(e)}"
        )
